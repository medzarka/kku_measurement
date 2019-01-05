from django.contrib import admin
from .models import SectionDocRequest
from .models import CourseDocRequest
from departments.models import Section
from .models import SectionAnalysisReport
from .models import CourseAnalysisReport
from periods.models import Semester


class SectionDocRequestAdmin(admin.ModelAdmin):
    list_display = (
        'section', 'created_date')
    list_filter = ('created_date',)

    def generate_state_list(self, request, queryset):
        import docx
        from django.http import HttpResponse

        _section_list = Section.objects.all()

        document = docx.Document()


        document.add_heading('List of Analysed Sections', 0)

        table = document.add_table(rows=1, cols=5)
        table.style = 'TableGrid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '#'
        hdr_cells[1].text = 'Section Number'
        hdr_cells[2].text = 'Course Code'
        hdr_cells[3].text = 'Course Name'
        hdr_cells[4].text = 'Teacher(s)'
        counter = 0
        for __section in _section_list:
            try:
                __tmp = SectionDocRequest.objects.get(section=__section)
                counter+=1
                row_cells = table.add_row().cells
                row_cells[0].text = str(counter)
                row_cells[1].text = str(__section.section_code)
                row_cells[2].text = __section.section_course.course_code
                row_cells[3].text = __section.section_course.course_name
                __tmp = ''
                for _teacher in __section.section_teachers.all():
                    __tmp = __tmp + _teacher.teacher_name_ar + '\n'
                row_cells[4].text = __tmp

            except SectionDocRequest.DoesNotExist:
                pass

        document.add_page_break()
        document.add_heading('List of Non Analysed Sections', 0)

        table = document.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        table.style = 'TableGrid'
        hdr_cells[0].text = '#'
        hdr_cells[1].text = 'Section Number'
        hdr_cells[2].text = 'Course Code'
        hdr_cells[3].text = 'Course Name'
        hdr_cells[4].text = 'Teacher(s)'
        counter = 0
        for __section in _section_list:
            try:
                __tmp = SectionDocRequest.objects.get(section=__section)
            except SectionDocRequest.DoesNotExist:
                counter += 1
                row_cells = table.add_row().cells
                row_cells[0].text = str(counter)
                row_cells[1].text = str(__section.section_code)
                row_cells[2].text = __section.section_course.course_code
                row_cells[3].text = __section.section_course.course_name
                __tmp = ''
                for _teacher in __section.section_teachers.all():
                    __tmp = __tmp + _teacher.teacher_name_ar + '\n'
                row_cells[4].text = __tmp
        # document.add_page_break()

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=download.docx'
        document.save(response)
        return response

    generate_state_list.short_description = "Generate State Report"

    actions = [generate_state_list, ]


class CourseDocRequestAdmin(admin.ModelAdmin):
    list_display = (
        'course', 'created_date')
    list_filter = ('created_date',)

class SectionAnalysisReportAdmin(admin.ModelAdmin):
    list_display = (
        'section', 'upload_date', 'file')
    list_filter = ('upload_date',)

    def generate_Section_Analysis_list(self, request, queryset):
        import docx
        import datetime
        from django.http import HttpResponse
        from docx.shared import Cm

        _section_list = Section.objects.all()
        __pending = []
        __ok = []
        for __section in _section_list:
            try:
                SectionAnalysisReport.objects.get(section=__section)
                __ok.append(__section)
            except SectionAnalysisReport.DoesNotExist:
                __pending.append(__section)

        document = docx.Document()

        d_date = datetime.datetime.now()
        reg_format_date = d_date.strftime("%Y-%m-%d %I:%M:%S %p")
        __actualSemester = Semester.objects.get(semester_isInUse=True)
        document.add_heading('List of Analyzed Sections', 0)
        p = document.add_paragraph('Measurement and Evaluation Unit')
        p.add_run(' -- ' + str(__actualSemester)).bold = True
        p = document.add_paragraph('Document generated at   ' + reg_format_date)

        document.add_heading('Analysed Section list', level=1)
        p = document.add_paragraph('In total,')
        p.add_run(' ' + str(len(__ok))).bold = True
        p.add_run(' analysed section(s).')
        table = document.add_table(rows=1, cols=4)
        table.style = 'TableGrid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Section'
        hdr_cells[1].text = 'Course'
        hdr_cells[2].text = 'Teacher(s)'
        hdr_cells[3].text = 'State'

        for __section in __ok:
            row_cells = table.add_row().cells
            row_cells[0].text = str(__section.section_code)
            row_cells[1].text = str(__section.section_course)
            __tmp = ''
            for _teacher in __section.section_teachers.all():
                __tmp = __tmp + _teacher.teacher_name_ar + '\n'
            row_cells[2].text = __tmp
            row_cells[3].text = SectionAnalysisReport.objects.get(section=__section).remarks




        document.add_page_break()
        document.add_heading('Pending Section list', level=1)
        p = document.add_paragraph('In total,')
        p.add_run(' ' + str(len(__pending))).bold = True
        p.add_run(' section(s) still pending analysis.')
        table = document.add_table(rows=1, cols=3)
        table.style = 'TableGrid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Section'
        hdr_cells[1].text = 'Course'
        hdr_cells[2].text = 'Teacher(s)'

        for __section in __pending:
            row_cells = table.add_row().cells
            row_cells[0].text = str(__section.section_code)
            row_cells[1].text = str(__section.section_course)
            __tmp = ''
            for _teacher in __section.section_teachers.all():
                __tmp = __tmp + _teacher.teacher_name_ar + '\n'
            row_cells[2].text = __tmp
        # document.add_page_break()

        sections = document.sections
        for section in sections:
            section.top_margin = Cm(1)
            section.bottom_margin = Cm(1)
            section.left_margin = Cm(1)
            section.right_margin = Cm(1)

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=download.docx'
        document.save(response)
        return response

    generate_Section_Analysis_list.short_description = "Generate Analyzed Section Reports list"
    actions = [generate_Section_Analysis_list, ]

class CourseAnalysisAdmin(admin.ModelAdmin):
    list_display = (
        'course', 'upload_date', 'file')
    list_filter = ('upload_date',)


admin.site.register(SectionDocRequest, SectionDocRequestAdmin)
admin.site.register(CourseDocRequest, CourseDocRequestAdmin)
admin.site.register(SectionAnalysisReport, SectionAnalysisReportAdmin)
admin.site.register(CourseAnalysisReport, CourseAnalysisAdmin)
